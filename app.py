
from datetime import date, datetime
from flask import Flask, render_template, redirect, url_for, request, flash, session, send_file, abort
import os
from config import Config
from models import db, Audit, NonConformita, AzioneCorrettiva, CustomField, CustomValue, ChecklistItem, Fornitore, InterventoManutenzione
from forms import NuovoAuditForm, SelezionaAuditForm, NuovaNCForm, NuovaAzioneForm, dynamic_fields_to_form
from sqlalchemy import func

# Auth
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps

login_manager = LoginManager()
login_manager.login_view = "login"  # se non autenticato -> /login

def create_app():
    app = Flask(__name__)
    app.config.from_object(Config)
    if not app.config.get("SECRET_KEY"):
        app.config["SECRET_KEY"] = "dev-secret-change-me"
    db.init_app(app)
    login_manager.init_app(app)

    # === BLUEPRINT: Formazione & Qualifica Fornitori ===
    from training import training_bp
    from suppliers import suppliers_bp

    # ====== MODEL UTENTE & RBAC ======
    class User(db.Model, UserMixin):
        __tablename__ = "users"
        id = db.Column(db.Integer, primary_key=True)
        name = db.Column(db.String(120), nullable=False)
        email = db.Column(db.String(255), unique=True, nullable=False)
        password_hash = db.Column(db.String(255), nullable=False)
        role = db.Column(db.String(20), nullable=False, default="junior")  # 'supervisor' | 'admin' | 'junior'
        requested_role = db.Column(db.String(20), nullable=True)           # richiesta di promozione

        # --- metodi utility ---
        def set_password(self, password: str):
            self.password_hash = generate_password_hash(password)

        def check_password(self, password: str) -> bool:
            return check_password_hash(self.password_hash, password)

        def is_supervisor(self) -> bool:
            return self.role == "supervisor"

        def is_admin(self) -> bool:
            return self.role == "admin" or self.is_supervisor()

    @login_manager.user_loader
    def load_user(user_id):
        try:
            return User.query.get(int(user_id))
        except Exception:
            return None

    

    # --- Supervisor utilities ---
    def _has_supervisor() -> bool:
        return db.session.query(User.id).filter(User.role == "supervisor").first() is not None

    def _supervisor_count() -> int:
            return db.session.query(User.id).filter(User.role == "supervisor").count()

# Bootstrap supervisor opzionale via env (non obbligatorio)
    def _bootstrap_supervisor():
        email = os.environ.get("SUPERVISOR_EMAIL")
        pwd   = os.environ.get("SUPERVISOR_PASSWORD")
        if not (email and pwd):
            return
        email = email.lower().strip()
        if not User.query.filter_by(email=email).first():
            su = User(name="Supervisor", email=email, role="supervisor")
            su.set_password(pwd)
            db.session.add(su)
            db.session.commit()

    # Whitelist endpoint consentiti ai junior
    JUNIOR_ALLOWLIST = {
        "index",                 # home/menu
        "list_audits",           # Elenco audit
        "actions_plan",          # Piano azioni correttive
        "richiedi_manutenzione", # Richiedi manutenzione (modulo)
        "login", "logout", "register", "static"
    }

    def roles_required(*roles):
        """Consenti accesso solo a determinati ruoli (supervisor sempre ammesso)."""
        def decorator(view):
            @wraps(view)
            def wrapped(*args, **kwargs):
                if not current_user.is_authenticated:
                    return login_manager.unauthorized()
                if current_user.is_supervisor():
                    return view(*args, **kwargs)
                if current_user.role not in roles:
                    abort(403)
                return view(*args, **kwargs)
            return wrapped
        return decorator

    @app.before_request
    def enforce_junior_limits():
        ep = request.endpoint or ""
        if ep.startswith("static"):
            return
        if not current_user.is_authenticated:
            return  # LoginManager gestirà il redirect
        if current_user.is_supervisor() or current_user.is_admin():
            return
        if current_user.role == "junior" and ep not in JUNIOR_ALLOWLIST:
            abort(403)

    with app.app_context():
        db.create_all()

        # Ensure special CustomField for NC notes exists
        note_cf = CustomField.query.filter_by(entity="nc", name="note").first()
        if not note_cf:
            note_cf = CustomField(entity="nc", name="note", label="Note", field_type="textarea", required=False)
            db.session.add(note_cf)
            db.session.commit()

        _bootstrap_supervisor()

    # ====== AUTH ROUTES ======
    @app.route("/register", methods=["GET", "POST"])
    def register():
        if request.method == "POST":
            name = (request.form.get("name") or "").strip()
            email = (request.form.get("email") or "").lower().strip()
            password = request.form.get("password") or ""
            desired_role = request.form.get("desired_role", "junior")  # 'junior' | 'admin' | 'supervisor'
  # 'junior' | 'admin'

            if not name or not email or not password:
                flash("Compila tutti i campi.", "warning")
                return redirect(url_for("register"))
            if User.query.filter_by(email=email).first():
                flash("Email già registrata.", "warning")
                return redirect(url_for("register"))

                        # Decide role: allow 'supervisor' only if none exists
            role_to_set = "junior"
            # Se è in corso la creazione del primo Supervisor, forza il ruolo
            if session.get('force_supervisor') and not _has_supervisor():
                role_to_set = "supervisor"
            if desired_role == "admin":
                role_to_set = "junior"
            # Se è in corso la creazione del primo Supervisor, forza il ruolo
            if session.get('force_supervisor') and not _has_supervisor():
                role_to_set = "supervisor"
            elif desired_role == "supervisor" and not _has_supervisor():
                role_to_set = "supervisor"
            u = User(name=name, email=email, role=role_to_set)
            u.set_password(password)
            if desired_role == "admin":
                u.requested_role = "admin"
            db.session.add(u)
            db.session.commit()
            if session.get('force_supervisor'):
                session.pop('force_supervisor', None)
            flash("Registrazione completata. Se hai chiesto 'Amministratore', il Supervisor approverà la richiesta.", "success")
            return redirect(url_for("login"))
        return render_template("register.html")

    @app.route("/login", methods=["GET", "POST"])
    def login():
        if request.method == "POST":
            email = (request.form.get("email") or "").lower().strip()
            password = request.form.get("password") or ""
            u = User.query.filter_by(email=email).first()
            if not u or not u.check_password(password):
                flash("Credenziali non valide.", "danger")
                return redirect(url_for("login"))
            login_user(u)
            return redirect(url_for("index"))
        return render_template("login.html")

    @app.route("/logout")
    @login_required
    def logout():
        logout_user()
        return redirect(url_for("login"))

    @app.route("/supervisor/reset", methods=["POST"])
    def supervisor_reset():
        can_reset = False
        dob = (request.form.get("dob") or "").strip()
        # Consenti reset se: data di nascita corretta OPPURE utente loggato supervisor
        if dob == "07021976":
            can_reset = True
        elif current_user.is_authenticated and current_user.is_supervisor():
            can_reset = True

        if not can_reset:
            abort(403)

        # Elimina tutti gli attuali supervisor
        supervisors = User.query.filter_by(role="supervisor").all()
        for s in supervisors:
            db.session.delete(s)
        db.session.commit()

        if current_user.is_authenticated:
            logout_user()

        flash("Supervisor azzerato. Registra ora un nuovo Supervisor per continuare.", "success")
        return redirect(url_for("new_supervisor"))



        # Pannello Supervisor per ruoli
    @app.route("/admin/users", methods=["GET", "POST"])
    @login_required
    @roles_required("admin")  # supervisor incluso
    def manage_users():
        if not current_user.is_supervisor():
            abort(403)  # solo supervisor modifica i ruoli
        if request.method == "POST":
            user_id = request.form.get("user_id")
            new_role = request.form.get("role")  # 'junior' | 'admin'
            u = User.query.get_or_404(user_id)
            if new_role not in ("junior", "admin"):
                flash("Ruolo non valido.", "warning")
                return redirect(url_for("manage_users"))
            u.role = new_role
            u.requested_role = None
            db.session.commit()
            flash(f"Ruolo aggiornato per {u.email}: {u.role}", "success")
            return redirect(url_for("manage_users"))
        users = User.query.order_by(User.id.asc()).all()
        return render_template("users_manage.html", users=users)

    # Elenco utenti (solo supervisore)
    @app.route("/admin/users-list", methods=["GET"])
    @login_required
    def users_list():
        if not current_user.is_supervisor():
            abort(403)
        users = User.query.order_by(User.id.asc()).all()
        return render_template("users_list.html", users=users, allow_supervisor_deletion=app.config.get("ALLOW_SUPERVISOR_DELETION", False))


    @app.post("/users/delete")
    @login_required
    def users_delete_post():
        # Solo Supervisor può eliminare utenti
        if not current_user.is_supervisor():
            abort(403)
        user_id = request.form.get("user_id")
        if not user_id or not str(user_id).isdigit():
            flash("Parametro user_id mancante o non valido.", "warning")
            return redirect(url_for("users_list"))
        u = User.query.get(int(user_id))
        if not u:
            flash("Utente non trovato.", "warning")
            return redirect(url_for("users_list"))

        # Gestione policy eliminazione Supervisor
        if u.role == "supervisor":
            allow = app.config.get("ALLOW_SUPERVISOR_DELETION", False)
            if not allow:
                flash("Eliminazione del Supervisor non consentita dalla policy.", "warning")
                return redirect(url_for("users_list"))
            # Non permettere di lasciare il sistema senza Supervisor
            if _supervisor_count() <= 1:
                flash("Non puoi eliminare l'ultimo Supervisor.", "warning")
                return redirect(url_for("users_list"))
            # Evita di eliminare te stessa/te stesso se resti senza privilegi
            self_delete = (u.id == current_user.id)
            db.session.delete(u)
            db.session.commit()
            flash("Supervisor eliminato correttamente.", "success")
            if self_delete:
                logout_user()
                return redirect(url_for("login"))
            return redirect(url_for("users_list"))

        # Elimina utenti non-supervisor
        db.session.delete(u)
        db.session.commit()
        flash("Utente eliminato correttamente.", "success")
        return redirect(url_for("users_list"))


    @app.post("/admin/users/<int:uid>/delete")
    @login_required
    def users_delete(uid):
        if not current_user.is_supervisor():
            abort(403)
        u = User.query.get_or_404(uid)
        if u.role == "supervisor":
            flash("Non puoi eliminare un utente con ruolo Supervisor.", "warning")
            return redirect(url_for("users_list"))
        db.session.delete(u); db.session.commit()
        flash("Utente eliminato.", "success")
        return redirect(url_for("users_list"))

    # ====== UTILS NC NOTES ======
    def _get_nc_note(nc_id):
        cf = CustomField.query.filter_by(entity="nc", name="note").first()
        if not cf:
            return ""
        cv = CustomValue.query.filter_by(field_id=cf.id, entity="nc", entity_id=nc_id).first()
        return cv.value if cv else ""

    def _set_nc_note(nc_id, text):
        cf = CustomField.query.filter_by(entity="nc", name="note").first()
        if not cf:
            cf = CustomField(entity="nc", name="note", label="Note", field_type="textarea", required=False)
            db.session.add(cf); db.session.commit()
        cv = CustomValue.query.filter_by(field_id=cf.id, entity="nc", entity_id=nc_id).first()
        if not cv:
            cv = CustomValue(field_id=cf.id, entity="nc", entity_id=nc_id, value=text or "")
            db.session.add(cv)
        else:
            cv.value = text or ""
        db.session.commit()

    # ====== ROUTES ======

    # NC: dettaglio con azioni
    @app.route("/ncs/<int:nc_id>", methods=["GET","POST"])
    @login_required
    def nc_detail(nc_id):
        nc = NonConformita.query.get_or_404(nc_id)
        if nc.stato == "Non applicabile":
            flash("Questa NC è marcata come Non applicabile e non è selezionabile.", "warning")
            return redirect(url_for("list_ncs", audit_id=nc.audit_id))
        return render_template("nc_detail.html", nc=nc)

    # NC: elimina (POST)
    @app.route("/ncs/<int:nc_id>/delete", methods=["POST"])
    @login_required
    def nc_delete(nc_id):
        nc = NonConformita.query.get_or_404(nc_id)
        aid = nc.audit_id
        CustomValue.query.filter_by(entity='nc', entity_id=nc_id).delete()
        db.session.delete(nc)
        db.session.commit()
        flash("Non conformità eliminata.", "success")
        return redirect(url_for("list_ncs", audit_id=aid))

    # NC: gestione (stato + note)
    @app.route("/ncs/<int:nc_id>/manage", methods=["GET","POST"])
    @login_required
    def nc_manage(nc_id):
        nc = NonConformita.query.get_or_404(nc_id)
        if request.method == "POST":
            nuovo_stato = request.form.get("stato")
            note = request.form.get("note","").strip()
            if nuovo_stato == "Risolta":
                nc.stato = "Chiusa"
            elif nuovo_stato == "Non applicabile":
                nc.stato = "Non applicabile"
            else:
                nc.stato = "Aperta"
            db.session.commit()
            _set_nc_note(nc.id, note)
            flash("NC aggiornata.", "success")
            return redirect(url_for("list_ncs", audit_id=nc.audit_id))
        note_val = _get_nc_note(nc.id)
        return render_template("nc_manage.html", nc=nc, note_val=note_val)

    
    

    @app.route("/supervisor/new", methods=["GET","POST"])
    def new_supervisor():
        # Se esiste già un supervisor, blocca
        if _has_supervisor():
            flash("Esiste già un Supervisor. Non è possibile crearne un altro direttamente.", "warning")
            return redirect(url_for("login"))
        if request.method == "POST":
            name = (request.form.get("name") or "").strip()
            email = (request.form.get("email") or "").lower().strip()
            password = request.form.get("password") or ""
            confirm = request.form.get("confirm") or ""
            if not name or not email or not password:
                flash("Compila tutti i campi.", "warning")
                return redirect(url_for("new_supervisor"))
            if password != confirm:
                flash("Le password non coincidono.", "warning")
                return redirect(url_for("new_supervisor"))
            # Check email già presente
            existing = User.query.filter_by(email=email).first()
            if existing:
                flash("Esiste già un utente con questa email. Usa un'altra email.", "warning")
                return redirect(url_for("new_supervisor"))
            # Crea utente supervisor
            u = User(name=name, email=email, role="supervisor")
            u.set_password(password)
            db.session.add(u)
            try:
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                flash("Errore nella creazione del Supervisor. Dettagli registrati nel log.", "danger")
                app.logger.exception("Errore creazione Supervisor: %s", e)
                return redirect(url_for("new_supervisor"))
            flash("Nuovo Supervisor creato con successo. Ora puoi accedere.", "success")
            return redirect(url_for("login"))
        return render_template("supervisor_new.html")

    @app.route("/")
    @login_required
    def index():
        return render_template("menu.html")

    # 1. Seleziona audit
    @app.route("/audits/select", methods=["GET","POST"])
    @login_required
    def select_audit():
        form = SelezionaAuditForm()
        audits = Audit.query.order_by(Audit.data_audit.desc()).all()
        form.audit_id.choices = [(a.id, f"{a.codice} — {a.titolo}") for a in audits]
        if form.validate_on_submit():
            session["current_audit_id"] = form.audit_id.data
            flash("Audit selezionato correttamente.", "success")
            return redirect(url_for("audit_detail", audit_id=form.audit_id.data))
        return render_template("select_audit.html", form=form)

    # 2. Nuovo audit
    @app.route("/audits/new", methods=["GET","POST"])
    @login_required
    def new_audit():
        cfields = CustomField.query.filter_by(entity="audit").all()
        DynamicAuditForm = dynamic_fields_to_form(NuovoAuditForm, cfields)
        form = DynamicAuditForm()
        if form.validate_on_submit():
            a = Audit(
                titolo=form.titolo.data,
                codice=form.codice.data,
                data_audit=form.data_audit.data,
                cliente=form.cliente.data,
                sede=form.sede.data,
                norma=form.norma.data,
                stato=form.stato.data,
            )
            db.session.add(a)
            db.session.flush()
            extra = {}
            for cf in cfields:
                val = getattr(form, f"cf_{cf.id}").data
                extra[str(cf.id)] = val.isoformat() if hasattr(val, "isoformat") else val
                db.session.add(CustomValue(field_id=cf.id, entity="audit", entity_id=a.id, value=str(extra[str(cf.id)])))
            a.extra = __import__("json").dumps(extra)
            db.session.commit()
            flash("Nuovo audit creato. Ora compila la Check list.", "success")
            return redirect(url_for("audit_checklist", audit_id=a.id))
        return render_template("new_audit.html", form=form)

    # 3. Elenco audit
    @app.route("/audits")
    @login_required
    def list_audits():
        q = request.args.get("q", "").strip()
        query = Audit.query
        if q:
            like = f"%%{q.lower()}%%"
            query = query.filter(func.lower(Audit.titolo).like(like) | func.lower(Audit.codice).like(like) | func.lower(Audit.cliente).like(like))
        audits = query.order_by(Audit.data_audit.desc()).all()
        return render_template("list_audits.html", audits=audits, q=q)

    # Dettaglio audit
    @app.route("/audits/<int:audit_id>")
    @login_required
    def audit_detail(audit_id):
        a = Audit.query.get_or_404(audit_id)
        ncs = NonConformita.query.filter_by(audit_id=a.id).order_by(NonConformita.data_apertura.desc()).all()
        actions = AzioneCorrettiva.query.filter_by(audit_id=a.id).order_by(AzioneCorrettiva.data_scadenza.asc()).all()
        return render_template("audit_detail.html", a=a, ncs=ncs, actions=actions)

    # 4. Elenco non conformità
    @app.route("/ncs", methods=["GET"])
    @login_required
    def list_ncs():
        audit_filter = request.args.get("audit_id", type=int)
        query = NonConformita.query
        if audit_filter:
            query = query.filter_by(audit_id=audit_filter)
        ncs = query.order_by(NonConformita.data_apertura.desc()).all()
        audits = Audit.query.order_by(Audit.data_audit.desc()).all()
        return render_template("list_ncs.html", ncs=ncs, audits=audits, audit_filter=audit_filter)

    # Nuova NC
    @app.route("/ncs/new", methods=["GET","POST"])
    @login_required
    def new_nc():
        audit_id = request.args.get("audit_id", type=int) or session.get("current_audit_id")
        if not audit_id:
            flash("Seleziona prima un audit.", "warning")
            return redirect(url_for("select_audit"))
        form = NuovaNCForm()
        return_url = request.args.get("return")
        prefill = request.args.get("prefill","")
        if request.method == "GET" and prefill:
            try:
                form.descrizione.data = prefill[:1000]
            except Exception:
                pass
        if form.validate_on_submit():
            nc = NonConformita(
                audit_id=audit_id,
                codice=form.codice.data,
                descrizione=form.descrizione.data,
                gravita=form.gravita.data,
                categoria=form.categoria.data,
                rilevata_da=form.rilevata_da.data,
                data_apertura=form.data_apertura.data or date.today(),
                stato=form.stato.data,
            )
            db.session.add(nc)
            db.session.commit()
            flash("Non conformità registrata.", "success")
            if request.form.get("return_url"):
                return redirect(request.form.get("return_url"))
            return redirect(url_for("list_ncs"))
        return render_template("new_nc.html", form=form, audit_id=audit_id, return_url=return_url)

    # 5. Piano delle azioni correttive
    @app.route("/actions", methods=["GET","POST"])
    @login_required
    def actions_plan():
        audit_id = request.args.get("audit_id", type=int) or session.get("current_audit_id")
        nc_id = request.args.get("nc_id", type=int)
        audits = Audit.query.order_by(Audit.data_audit.desc()).all()
        selected_audit = Audit.query.get(audit_id) if audit_id else None
        selected_nc = NonConformita.query.get(nc_id) if nc_id else None
        if request.method == "POST":
            form = NuovaAzioneForm()
            if form.validate_on_submit():
                action = AzioneCorrettiva(
                    audit_id=audit_id or (selected_nc.audit_id if selected_nc else None),
                    nc_id=selected_nc.id if selected_nc else request.args.get("nc_id", type=int),
                    azione=form.azione.data,
                    responsabile=form.responsabile.data,
                    data_scadenza=form.data_scadenza.data,
                    stato=form.stato.data,
                    efficacia=form.efficacia.data,
                )
                db.session.add(action)
                db.session.commit()
                flash("Azione correttiva aggiunta.", "success")
                return redirect(url_for("actions_plan", audit_id=(audit_id or (selected_nc.audit_id if selected_nc else None)), nc_id=(selected_nc.id if selected_nc else None)))
        form = NuovaAzioneForm()
        actions = AzioneCorrettiva.query
        if audit_id or (selected_nc and selected_nc.audit_id):
            actions = actions.filter_by(audit_id=(audit_id or selected_nc.audit_id))
        actions = actions.order_by(AzioneCorrettiva.data_scadenza.asc().nullslast()).all()
        ncs = NonConformita.query.filter_by(audit_id=(audit_id or (selected_nc.audit_id if selected_nc else None))).all() if (audit_id or selected_nc) else []
        return render_template("actions_plan.html", audits=audits, selected_audit=selected_audit, actions=actions, ncs=ncs, form=form, selected_nc=selected_nc)

    # --- Configurazione campi dinamici ---
    @app.route("/config/campi-audit", methods=["GET","POST"])
    @login_required
    def config_campi_audit():
        if request.method == "POST":
            name = request.form.get("name","").strip()
            label = request.form.get("label","").strip()
            field_type = request.form.get("field_type","text")
            required = True if request.form.get("required") == "on" else False
            options = request.form.get("options","").strip()
            cf = CustomField(entity="audit", name=name, label=label, field_type=field_type, options=options or None, required=required)
            db.session.add(cf)
            db.session.commit()
            flash("Campo personalizzato creato.", "success")
            return redirect(url_for("config_campi_audit"))
        fields = CustomField.query.filter_by(entity="audit").all()
        return render_template("config_campi_audit.html", fields=fields)

    @app.context_processor
    def inject_now():
        return {"now": datetime.utcnow()}

    # ---- CHECKLIST ----
    @app.route("/audits/<int:audit_id>/checklist", methods=["GET","POST"])
    @login_required
    def audit_checklist(audit_id):
        a = Audit.query.get_or_404(audit_id)
        # Import from DOCX
        if request.method == "POST" and request.form.get("action") == "import":
            file = request.files.get("docx")
            if not file or not file.filename.lower().endswith(".docx"):
                flash("Carica un file .docx valido.", "warning")
                return redirect(url_for("audit_checklist", audit_id=audit_id))
            try:
                from docx import Document
                import io, re
                doc = Document(io.BytesIO(file.read()))
                def clean_text(t):
                    t = t.strip()
                    t = re.sub(r"^(\s*[-*•]|\s*\d+[\.)])\s*", "", t)
                    return t.strip()
                items = []
                if len(doc.tables) > 0:
                    for table in doc.tables:
                        cols = len(table.columns)
                        for r_idx, row in enumerate(table.rows):
                            cells = [c.text.strip() for c in row.cells]
                            if r_idx == 0 and cells and cells[0] and (("voce" in cells[0].lower()) or ("descr" in cells[0].lower())):
                                continue
                            if not cells or (cells[0].strip() == ""):
                                continue
                            descr = clean_text(cells[0])
                            esito = cells[1].strip() if cols >= 2 else ""
                            note = cells[2].strip() if cols >= 3 else ""
                            items.append((descr, esito, note))
                else:
                    for p in doc.paragraphs:
                        style = (getattr(p.style, "name", "") or "").lower()
                        if "heading" in style or "titolo" in style:
                            continue
                        txt = clean_text(p.text)
                        if txt:
                            items.append((txt, "", ""))
                def map_esito(v):
                    s = (v or "").strip().lower()
                    if s in ("conforme","ok","si","sì","si'","yes"): return "Conforme"
                    if s in ("non conforme","nc","ko","no"): return "Non conforme"
                    if s in ("non applicabile","n/a","na"): return "Non applicabile"
                    return "Aperta"
                max_ord = db.session.query(func.max(ChecklistItem.ordine)).filter_by(audit_id=audit_id).scalar() or 0
                added = 0
                for descr, esito, note in items[:1000]:
                    if not descr:
                        continue
                    added += 1
                    max_ord += 1
                    item = ChecklistItem(audit_id=audit_id, descrizione=descr, ordine=max_ord, esito=map_esito(esito), note=(note or None))
                    db.session.add(item)
                db.session.commit()
                flash(f"Import completato: {added} righe aggiunte dalla check list.", "success")
            except Exception as e:
                flash(f"Errore durante l'import: {e}", "warning")
            return redirect(url_for("audit_checklist", audit_id=audit_id))
        # Add new row
        if request.method == "POST" and request.form.get("action") == "add":
            descr = request.form.get("descrizione","").strip()
            if descr:
                max_ord = db.session.query(func.max(ChecklistItem.ordine)).filter_by(audit_id=audit_id).scalar() or 0
                item = ChecklistItem(audit_id=audit_id, descrizione=descr, ordine=max_ord+1)
                db.session.add(item)
                db.session.commit()
                flash("Riga di check list aggiunta.", "success")
            return redirect(url_for("audit_checklist", audit_id=audit_id))

        # Save statuses
        if request.method == "POST" and request.form.get("action") == "save":
            for key, val in request.form.items():
                if key.startswith("esito_"):
                    item_id = int(key.split("_",1)[1])
                    item = ChecklistItem.query.get(item_id)
                    if item and item.audit_id == a.id:
                        item.esito = val
                if key.startswith("note_"):
                    item_id = int(key.split("_",1)[1])
                    item = ChecklistItem.query.get(item_id)
                    if item and item.audit_id == a.id:
                        item.note = request.form.get(key,"").strip()
            db.session.commit()
            flash("Check list salvata.", "success")
            return redirect(url_for("audit_checklist", audit_id=audit_id))

        items = ChecklistItem.query.filter_by(audit_id=audit_id).order_by(ChecklistItem.ordine.asc(), ChecklistItem.id.asc()).all()
        return render_template("checklist.html", a=a, items=items)

    @app.route("/audits/<int:audit_id>/checklist/export", methods=["GET"])
    @login_required
    def export_checklist(audit_id):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from xml.sax.saxutils import escape

        a = Audit.query.get_or_404(audit_id)
        items = (ChecklistItem.query
                 .filter_by(audit_id=audit_id)
                 .order_by(ChecklistItem.ordine.asc(), ChecklistItem.id.asc())
                 .all())

        pdf_path = os.path.join(app.instance_path, f"checklist_audit_{a.id}.pdf")
        os.makedirs(app.instance_path, exist_ok=True)

        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            leftMargin=18*mm, rightMargin=18*mm,
            topMargin=20*mm, bottomMargin=18*mm
        )

        styles = getSampleStyleSheet()
        # wordWrap='CJK' per capo anche su parole lunghe
        cell = ParagraphStyle('cell', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12, spaceBefore=0, spaceAfter=0, wordWrap='CJK')
        head = ParagraphStyle('head', parent=styles['Heading5'], fontName='Helvetica-Bold', fontSize=10, leading=12, spaceBefore=0, spaceAfter=4)

        story = []
        story.append(Paragraph(f"Check list — Audit {a.codice}", styles["Title"]))
        meta = f"{escape(a.titolo or '—')} • {escape(a.cliente or '—')} • {escape(a.sede or '—')} • {a.data_audit.strftime('%d/%m/%Y')}"
        story.append(Paragraph(meta, styles["Normal"]))
        story.append(Spacer(1, 6))

        content_w = doc.width
        w_num, w_esito, w_note = 10*mm, 30*mm, 55*mm
        w_voce = content_w - (w_num + w_esito + w_note)
        if w_voce < 40*mm:
            deficit = (40*mm) - w_voce
            w_voce = 40*mm
            w_note = max(30*mm, w_note - deficit)

        data = [[Paragraph("#", head), Paragraph("Voce di verifica", head), Paragraph("Esito", head), Paragraph("Note", head)]]
        for idx, it in enumerate(items, start=1):
            voce = Paragraph(escape(it.descrizione or ""), cell)
            esito = Paragraph(escape(it.esito or ""), cell)
            note = Paragraph(escape(it.note or ""), cell)
            data.append([str(idx), voce, esito, note])

        tbl = Table(data, colWidths=[w_num, w_voce, w_esito, w_note], repeatRows=1, hAlign='LEFT')
        tbl.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.whitesmoke),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 4),
            ("RIGHTPADDING", (0,0), (-1,-1), 4),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        story.append(tbl)
        doc.build(story)
        return send_file(pdf_path, as_attachment=True, download_name=f"Checklist_Audit_{a.codice}.pdf")

    @app.route("/audits/<int:audit_id>/delete", methods=["POST"])
    @login_required
    def delete_audit(audit_id):
        a = Audit.query.get_or_404(audit_id)
        db.session.delete(a)
        db.session.commit()
        flash("Audit eliminato.", "success")
        return redirect(url_for("list_audits"))

    # === Richieste di manutenzione ===
    @app.route("/manutenzione", methods=["GET", "POST"])
    @login_required
    def richiedi_manutenzione():
        # Creazione PDF + salvataggio richiesta
        if request.method == "POST" and request.form.get("azione") == "genera_pdf":
            tipo = request.form.get("tipo") or ""
            descrizione = request.form.get("descrizione") or ""
            richiedente = request.form.get("richiedente") or ""
            centro = request.form.get("centro") or ""
            data_intervento = request.form.get("data_intervento") or ""

            fornitore = None
            destinatario = None
            tipo_label = ""
            if tipo == "elettrica":
                tipo_label = "Manutenzione elettrica"; fornitore = request.form.get("fornitore") or ""
            elif tipo == "meccanica":
                tipo_label = "Manutenzione meccanica"; fornitore = request.form.get("fornitore") or ""
            elif tipo == "servizi":
                tipo_label = "Servizi"; destinatario = request.form.get("destinatario") or ""

            html = render_template(
                "Richiesta di intervento.html",
                tipo_label=tipo_label, descrizione=descrizione, richiedente=richiedente,
                centro=centro, fornitore=fornitore or None, destinatario=destinatario or None,
                data_intervento=data_intervento, data_generazione=datetime.now().strftime("%d/%m/%Y %H:%M")
            )

            from xhtml2pdf import pisa
            from io import BytesIO
            pdf_io = BytesIO()
            pisa.CreatePDF(src=html, dest=pdf_io)
            pdf_io.seek(0)

            # Salvataggio nel DB
            try:
                try:
                    from datetime import datetime as _dt
                    di = _dt.strptime(data_intervento, "%Y-%m-%d").date() if data_intervento else None
                except Exception:
                    di = None
                rec = InterventoManutenzione(
                    tipo=tipo, descrizione=descrizione, richiedente=richiedente, centro=centro,
                    fornitore=fornitore or None, destinatario=destinatario or None,
                    data_intervento=di, status="Non assegnato"
                )
                db.session.add(rec); db.session.commit()
            except Exception as e:
                db.session.rollback()
                flash(f"Errore nel salvataggio della richiesta: {e}", "danger")

            return send_file(pdf_io, as_attachment=True, download_name="Richiesta_intervento.pdf", mimetype="application/pdf")

        # Salvataggio nuovo fornitore
        if request.method == "POST" and request.form.get("azione") == "salva_fornitore":
            nome_f = (request.form.get("nuovo_nome_fornitore") or "").strip()
            tel_f = (request.form.get("nuovo_telefono") or "").strip()
            contatto_f = (request.form.get("nuovo_contatto") or "").strip()
            tipologia_f = (request.form.get("nuova_tipologia") or "").strip()
            if not nome_f or tipologia_f not in ("elettrica", "meccanica", "servizi"):
                flash("Compila almeno Nome fornitore e Tipologia corretti.", "warning")
            else:
                f = Fornitore(nome=nome_f, telefono=tel_f, contatto=contatto_f, tipologia=tipologia_f)
                db.session.add(f); db.session.commit()
                flash(f"Fornitore '{nome_f}' salvato.", "success")
            tipo_sel = tipologia_f
        else:
            tipo_sel = request.args.get("tipo", "")

        fornitori = []
        if tipo_sel in ("elettrica", "meccanica"):
            fornitori = [x.nome for x in Fornitore.query.filter_by(tipologia=tipo_sel).order_by(Fornitore.nome).all()]
        return render_template("manutenzione.html", fornitori=fornitori)

    # Elenco interventi (ENDPOINT RICHIESTO)
    @app.route("/manutenzione/elenco", methods=["GET"])
    @login_required
    def elenco_manutenzioni():
        # compat con modelli senza created_at
        order_col = getattr(InterventoManutenzione, "created_at", None)
        q = InterventoManutenzione.query
        if order_col is not None:
            rows = q.order_by(order_col.desc()).all()
        else:
            rows = q.order_by(InterventoManutenzione.id.desc()).all()
        return render_template("manutenzione_list.html", rows=rows)

    @app.route("/manutenzione/<int:mid>/delete", methods=["POST"])
    @login_required
    def delete_manutenzione(mid):
        r = InterventoManutenzione.query.get_or_404(mid)
        db.session.delete(r); db.session.commit()
        flash("Intervento eliminato.", "success")
        return redirect(url_for("elenco_manutenzioni"))

    @app.route("/manutenzione/<int:mid>/status", methods=["POST"])
    @login_required
    def update_manutenzione_status(mid):
        r = InterventoManutenzione.query.get_or_404(mid)
        status = (request.form.get("status") or "").strip()
        allowed = {"Eseguito", "Assegnato", "Non assegnato"}
        if status not in allowed:
            flash("Stato non valido.", "warning")
        else:
            r.status = status
            db.session.commit()
            flash("Stato aggiornato.", "success")
        # se la lista è in un blueprint, qualifica:
        # return redirect(url_for("manutenzione.elenco"))
        return redirect(url_for("elenco_manutenzioni"))

    # Registrazione blueprint QUI, dentro la factory
    app.register_blueprint(training_bp, url_prefix="/formazione")
    app.register_blueprint(suppliers_bp, url_prefix="/qualifica-fornitori")

    return app

app = create_app()

if __name__ == '__main__':
    import webbrowser
    from threading import Timer
    def open_browser():
        webbrowser.open_new("http://127.0.0.1:8080/login")
    Timer(1.0, open_browser).start()
    app.run(host='0.0.0.0', port=8080, use_reloader=False)
